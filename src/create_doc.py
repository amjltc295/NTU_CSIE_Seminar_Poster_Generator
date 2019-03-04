import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from logging_config import logger


def create_doc(
    name, title, speecher, time, location,
    output_dir=os.path.join(os.path.dirname(os.path.dirname(os.path.realpath(__file__))), 'outputs')
):
    document = Document()

    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.right_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('國立台灣大學電機資訊學院演講公告\n')
    font = run.font
    font.name = 'Liberation Serif'
    font.size = Pt(43.5)
    font.bold = True
    font.underline = True

    style = document.styles['Normal']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.667)
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(36)
    p.add_run(f"講題：{title}\n")
    p.add_run(f"講員：{speecher}\n")
    p.add_run(f"時間：{time}\n")
    p.add_run(f"地點：{location}")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("【專題討論演講】").bold = True

    document.add_page_break()

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        logger.info(f"Made {output_dir}")
    output_filename = os.path.join(output_dir, f"{name}.docx")
    logger.info(f"Createing doc {output_filename}")
    document.save(output_filename)


if __name__ == "__main__":
    create_doc('Test', 'test title', 'me', '20191313', 'R506')
